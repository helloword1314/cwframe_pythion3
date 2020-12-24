# -*- coding: utf-8 -*-
import os
import sys
import docx
from docx import Document
# 如果不行，加上之前的基础信息
try :
    basic_python
except NameError:
    sys.path.append('\\'.join(os.path.realpath(__file__).split('\\')[:-2]))
    import basic_python.basic_function_python as basic_function_python


def Get_Paragraph_List_In_Doc(input_docx):
    '''得到文档中的段落数组
    Args:  
    '''
    arrParas=[]
    doc = Document(input_docx)
    for para in doc.paragraphs:
        arrParas.append(para.text)
    return arrParas

def Create_Document():
    '''创建文档
    Args:  
    '''
    document=Document()
    return document
 
def Open_Document(input_docx):
    '''创建文档
    Args:  
    '''
    document=Document(input_docx)
    return document
	
def Save_Document(input_docx,document):
    '''保存文档
    Args:  
    '''
    document.save(input_docx)
	
def Set_Document_Style(document):
	'''创建文档
	Args:  
	'''
	document.styles['Normal'].font.name = u'宋体'
	document.styles['Normal']._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'宋体')
	document.styles['Normal'].font.size = docx.shared.Pt(14)
	document.styles['Normal'].font.color.rgb = docx.shared.RGBColor(0,0,0)

def Add_Heading_And_Set_Style(str_head,int_level,document):
	'''创建文档
	Args:  
	'''
	Head = document.add_heading("",level=int_level)
	run  = Head.add_run(str_head)
	run.font.name=u'Cambria'
	fontsize=12
	if int_level==1:
		fontsize=16
	elif int_level==2:
		fontsize=14
	else:
		fontsize=12
	run.font.size=docx.shared.Pt(fontsize)
	run._element.rPr.rFonts.set(docx.oxml.ns.qn('w:eastAsia'), u'Cambria')
	run.font.color.rgb = docx.shared.RGBColor(0,0,0)
	
def Explain_Something_About_Font():	
	''' 
	字号‘八号’对应磅值5 字号‘七号’对应磅值5.5 字号‘小六’对应磅值6.5字号‘六号’对应磅值7.5
	字号‘小五’对应磅值9字号‘五号’对应磅值10.5字号‘小四’对应磅值12字号‘四号’对应磅值14
	字号‘小三’对应磅值15字号‘三号’对应磅值16字号‘小二’对应磅值18字号‘二号’对应磅值22
	字号‘小一’对应磅值24字号‘一号’对应磅值26字号‘小初’对应磅值36字号‘初号’对应磅值42
	'''
	pass
 
def Add_Paragraph_String_To_Document(strPara,document_name):
    '''创建文档
    Args:  
    '''
    document=Document(document_name)
    document.add_paragraph(strPara).add_run().font.size=Pt(4)
    document.save(document_name)

def Add_Page_Break(document):
	'''创建文档
    Args:  
    '''	
	document.add_page_break()

def Add_Picture_And_Set_Size(input_pic,size,document):
	'''创建文档
    Args:  
    '''	
	document.add_picture(input_pic, width=docx.shared.Inches(size), height=docx.shared.Inches(size))

def Add_Table_Use_Rows_And_Cols(int_rows,int_cols,document):
	'''创建文档
    Args:  
    '''	
	document.add_table(rows=int_rows, cols=int_cols)
	
def Say_Something():
    '''修改
    Args:
        
    '''
    print("曹维，人生真是无聊啊！")
	
	
if __name__=="__main__":
    Say_Something()
